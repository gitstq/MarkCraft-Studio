"""
DOCX导出器
将Markdown转换为Word文档
使用python-docx生成格式化的DOCX文件
"""

import os
import re
from markcraft.core.engine import MarkCraftEngine, ExportConfig


class DOCXExporter:
    """DOCX格式导出器"""

    def export(
        self,
        markdown_text: str,
        output_path: str,
        config: ExportConfig,
        engine: MarkCraftEngine
    ) -> str:
        """
        导出为DOCX文件

        Args:
            markdown_text: Markdown原始文本
            output_path: 输出文件路径
            config: 导出配置
            engine: MarkCraft引擎实例

        Returns:
            str: 输出文件的完整路径
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            raise ImportError(
                "DOCX导出需要安装python-docx库。请运行: pip install python-docx"
            )

        # 解析Markdown
        result = engine.convert(markdown_text, config)

        # 创建Word文档
        doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(11)

        # 设置页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

        # 添加标题
        title = config.title
        if result.metadata.get('title'):
            title = result.metadata['title']

        doc.add_heading(title, level=0)

        if config.author or result.metadata.get('author'):
            author = config.author or result.metadata.get('author', '')
            p = doc.add_paragraph()
            run = p.add_run(f"作者：{author}")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # 添加统计信息
        stats = doc.add_paragraph()
        stats_run = stats.add_run(
            f"字数：{result.word_count} · 预计阅读：{result.reading_time}分钟"
        )
        stats_run.font.size = Pt(9)
        stats_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # 按行处理Markdown内容
        lines = markdown_text.split('\n')
        i = 0
        in_code_block = False
        code_lines = []
        code_lang = ''

        while i < len(lines):
            line = lines[i]

            # 跳过YAML前置数据
            if i == 0 and line.strip() == '---':
                while i < len(lines) and lines[i].strip() != '---':
                    i += 1
                i += 1
                continue

            # 代码块处理
            if line.strip().startswith('```'):
                if not in_code_block:
                    in_code_block = True
                    code_lang = line.strip()[3:].strip()
                    code_lines = []
                else:
                    in_code_block = False
                    # 添加代码块
                    code_text = '\n'.join(code_lines)
                    p = doc.add_paragraph(code_text)
                    p.style = doc.styles['Normal']
                    for run in p.runs:
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                    # 添加代码块背景（通过段落底纹）
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'F5F5F5')
                    shading.set(qn('w:val'), 'clear')
                    p.paragraph_format.element.get_or_add_pPr().append(shading)
                i += 1
                continue

            if in_code_block:
                code_lines.append(line)
                i += 1
                continue

            # 标题处理
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                doc.add_heading(text, level=min(level, 4))
                i += 1
                continue

            # 水平线
            if re.match(r'^(---|\*\*\*|___)\s*$', line.strip()):
                doc.add_paragraph('─' * 50)
                i += 1
                continue

            # 引用块
            if line.strip().startswith('> '):
                quote_text = line.strip()[2:]
                p = doc.add_paragraph()
                run = p.add_run(quote_text)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                i += 1
                continue

            # 无序列表
            list_match = re.match(r'^(\s*)[-*+]\s+(.+)$', line)
            if list_match:
                text = list_match.group(2)
                p = doc.add_paragraph(text, style='List Bullet')
                i += 1
                continue

            # 有序列表
            ol_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
            if ol_match:
                text = ol_match.group(2)
                p = doc.add_paragraph(text, style='List Number')
                i += 1
                continue

            # 空行跳过
            if not line.strip():
                i += 1
                continue

            # 普通段落
            text = self._process_inline_formatting(line.strip(), doc)
            if text:
                doc.add_paragraph(text)

            i += 1

        # 确保输出目录存在
        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)

        doc.save(output_path)
        return os.path.abspath(output_path)

    @staticmethod
    def _process_inline_formatting(text: str, doc) -> str:
        """处理行内格式（粗体、斜体、代码等）"""
        # 简单处理：返回纯文本（DOCX的行内格式需要更复杂的处理）
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'`(.+?)`', r'\1', text)
        text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
        return text
